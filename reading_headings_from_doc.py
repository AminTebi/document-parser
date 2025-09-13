from docx import Document


# Load the document
doc_path = "PASTE_YOUR_FILE_PATH_HERE"

def extract_hierarchical_content(doc_path):
    """
    Extract content in hierarchical structure:
    Heading 1 -> Heading 2 -> Content (string or list of strings)
    """
    doc = Document(doc_path)
    result = {}
    current_heading1 = None
    current_heading2 = None
    current_content = []
    current_content_type = None  # 'single' or 'list'

    def save_current_content():
        """Helper function to save current content with proper formatting"""
        if current_heading1 and current_heading2:
            if current_heading1 not in result:
                result[current_heading1] = {}
            
            # Determine content format
            if current_content_type == 'list' and len(current_content) > 1:
                result[current_heading1][current_heading2] = current_content
            elif current_content_type == 'list' and len(current_content) == 1:
                result[current_heading1][current_heading2] = current_content[0]
            elif current_content_type == 'single':
                # Check if single content should be split into list
                content_text = current_content[0] if current_content else ""
                if '\n' in content_text and len(content_text.split('\n')) > 1:
                    # Split into list of strings
                    lines = [line.strip() for line in content_text.split('\n') if line.strip()]
                    if len(lines) > 1:
                        result[current_heading1][current_heading2] = lines
                    else:
                        result[current_heading1][current_heading2] = content_text
                else:
                    result[current_heading1][current_heading2] = content_text

    def is_bullet_point(text):
        """Check if text appears to be a bullet point"""
        # Check for common bullet point patterns
        bullet_patterns = [
            text.startswith('•'),
            text.startswith('◦'),
            text.startswith('▪'),
            text.startswith('▫'),
            text.startswith('-'),
            text.startswith('*'),
            text.startswith('+'),
            # Check for numbered lists
            any(text.startswith(f'{i}.') for i in range(1, 20)),
            any(text.startswith(f'{i})') for i in range(1, 20)),
        ]
        return any(bullet_patterns)

    def split_bullet_content(text):
        """Split text that contains multiple bullet points separated by newlines"""
        lines = text.split('\n')
        bullet_items = []
        current_item = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if is_bullet_point(line):
                if current_item:
                    bullet_items.append(current_item.strip())
                current_item = line
            else:
                if current_item:
                    current_item += "\n" + line
                else:
                    current_item = line
        
        if current_item:
            bullet_items.append(current_item.strip())
            
        return bullet_items

    for p in doc.paragraphs:
        style = p.style.name
        text = p.text.strip()

        if not text:
            continue  # skip empty lines

        # Handle Heading 1
        if style == "Heading 1":
            # Save previous heading 2 if exists
            save_current_content()
            
            # Start new Heading 1
            current_heading1 = text
            current_heading2 = None
            current_content = []
            current_content_type = None

        # Handle Heading 2
        elif style == "Heading 2":
            # Save previous heading 2 content if exists
            save_current_content()
            
            # Start new Heading 2
            current_heading2 = text
            current_content = []
            current_content_type = None

        # Handle bullet points (detect by content pattern)
        elif is_bullet_point(text) or (current_content_type == 'list' and '\n' in text):
            if current_content_type is None:
                current_content_type = 'list'
                current_content = []
            elif current_content_type == 'single':
                # Convert single content to list
                current_content = [current_content[0]] if current_content else []
                current_content_type = 'list'
            
            # If text contains multiple bullet points, split them
            if '\n' in text and any(is_bullet_point(line.strip()) for line in text.split('\n') if line.strip()):
                bullet_items = split_bullet_content(text)
                current_content.extend(bullet_items)
            else:
                current_content.append(text)

        # Handle normal text (paragraph content)
        else:
            if current_content_type is None:
                current_content_type = 'single'
                current_content = [text]
            elif current_content_type == 'single':
                # Append to existing single content
                if current_content:
                    current_content[0] += "\n" + text
                else:
                    current_content = [text]
            elif current_content_type == 'list':
                # Add to list as separate item
                current_content.append(text)

    # Save the last heading 2 if exists
    save_current_content()

    return result


def extract_heading2_content(doc_path):
    """
    Legacy function - kept for backward compatibility
    """
    doc = Document(doc_path)
    result = {}
    current_heading = None
    buffer = []

    for p in doc.paragraphs:
        style = p.style.name
        text = p.text.strip()

        if not text:
            continue  # skip empty lines

        # When we hit a Heading 2, save the previous one and start fresh
        if style == "Heading 2":
            if current_heading:  # save previous section
                result[current_heading] = buffer
            current_heading = text
            buffer = []

        # Handle bullet points (usually "List Bullet" style)
        elif style.startswith("List"):
            if not isinstance(buffer, list):
                buffer = []
            buffer.append(text)

        # Normal text (paragraph content)
        else:
            if isinstance(buffer, list):
                # already collecting bullet points
                buffer.append(text)
            else:
                if buffer:
                    buffer += "\n" + text
                else:
                    buffer = text

    # Save the last heading
    if current_heading:
        result[current_heading] = buffer

    return result






# Extract hierarchical content (Heading 1 -> Heading 2 -> Content)
hierarchical_result = extract_hierarchical_content(doc_path)

# Print the hierarchical structure
print("Hierarchical Structure:")
print("=" * 50)
for heading1, heading2_dict in hierarchical_result.items():
    print(f"\n{heading1}:")
    for heading2, content in heading2_dict.items():
        print(f"  {heading2}: {content}")

# Also print as JSON for better visualization
import json
print("\n" + "=" * 50)
print("JSON Format:")
print(json.dumps(hierarchical_result, indent=2, ensure_ascii=False))